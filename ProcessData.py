import os
import json
from docx import Document

# Define folder paths
data_folder = "data_files"  # Folder containing JSON data files
output_folder = "output_documents"  # Folder for saving Word documents

# Create output folder if it doesn't exist
os.makedirs(output_folder, exist_ok=True)

# Iterate through all JSON data files in the folder
for file_name in os.listdir(data_folder):
    if file_name.endswith('.json'):  # Check for JSON files
        # Extract job class name (file name without extension)
        job_class = os.path.splitext(file_name)[0]
        
        # Load the JSON data file
        file_path = os.path.join(data_folder, file_name)
        with open(file_path, 'r') as file:
            tasks = json.load(file)  # Load the JSON array
        
        # Create a new Word document
        doc = Document()
        
        # Add title and job class
        doc.add_heading("AI-Assisted Task Recommendations", level=0)
        subtitle = doc.add_paragraph(f"Job Class: {job_class}")
        subtitle.style = 'Subtitle'  # Apply the Word 'Subtitle' style
        
        doc.add_paragraph()  # Add a blank line for spacing
        
        # Add tasks to the document
        for task in tasks:  # Iterate through the array
            doc.add_heading(f"Task: {task['Task Name']}", level=1)
            doc.add_heading(f"Microsoft Product: {task['Microsoft Product']}", level=2)
            doc.add_paragraph(task['Description'])
            doc.add_paragraph("---")  # Add a separator for tasks
        
        # Save the document with the job class name
        output_path = os.path.join(output_folder, f"{job_class}.docx")
        doc.save(output_path)
        
        print(f"Generated document for job class: {job_class}")

print("All documents have been generated.")
